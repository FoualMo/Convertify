import os
from docx import Document
from fpdf import FPDF
from PIL import Image
import pdfplumber
import markdown
import html2text

class FileConverter:
    SUPPORTED_FORMATS = ['docx', 'pdf', 'txt', 'md', 'html', 'jpg', 'jpeg', 'png', 'bmp', 'webp']
    OUTPUT_FOLDER = "static/converted"

    @classmethod
    def ensure_output_folder(cls):
        os.makedirs(cls.OUTPUT_FOLDER, exist_ok=True)

    # --- Extraction de texte ---
    @staticmethod
    def pdf_to_text(input_path):
        try:
            with pdfplumber.open(input_path) as pdf:
                return "\n".join(page.extract_text() or "" for page in pdf.pages)
        except Exception as e:
            raise ValueError(f"Erreur lecture PDF : {e}")

    @staticmethod
    def docx_to_text(input_path):
        try:
            doc = Document(input_path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            raise ValueError(f"Erreur lecture DOCX : {e}")

    @staticmethod
    def txt_to_text(input_path):
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Erreur lecture TXT : {e}")

    @staticmethod
    def md_to_text(input_path):
        return FileConverter.txt_to_text(input_path)

    @staticmethod
    def html_to_text(input_path):
        try:
            with open(input_path, "r", encoding="utf-8") as f:
                html_content = f.read()
            return html2text.html2text(html_content)
        except Exception as e:
            raise ValueError(f"Erreur lecture HTML : {e}")

    # --- Conversion vers PDF ---
    @staticmethod
    def text_to_pdf(text, output_path, font_size=12):
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font("Arial", size=font_size)
        for line in text.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf.output(output_path)

    @staticmethod
    def docx_to_pdf(input_path, output_path):
        text = FileConverter.docx_to_text(input_path)
        FileConverter.text_to_pdf(text, output_path)

    @staticmethod
    def image_to_pdf(input_path, output_path):
        try:
            img = Image.open(input_path).convert("RGB")
            img.save(output_path)
        except Exception as e:
            raise ValueError(f"Erreur conversion image → PDF : {e}")

    # --- Conversion finale ---
    @classmethod
    def convert(cls, input_path, output_format):
        cls.ensure_output_folder()

        name, ext = os.path.splitext(os.path.basename(input_path))
        ext = ext.lower()
        output_format = output_format.lower()
        output_path = os.path.join(cls.OUTPUT_FOLDER, f"{name}.{output_format}")

        if ext not in ['.pdf', '.docx', '.txt', '.md', '.html', '.jpg', '.jpeg', '.png', '.bmp', '.webp']:
            raise ValueError(f"Format d'entrée non supporté : {ext}")

        # --- Cas images ---
        if ext in [".jpg", ".jpeg", ".png", ".bmp", ".webp"]:
            if output_format in ["jpg", "jpeg", "png", "bmp", "webp"]:
                img = Image.open(input_path)
                img.save(output_path)
                return output_path
            elif output_format == "pdf":
                cls.image_to_pdf(input_path, output_path)
                return output_path
            else:
                raise ValueError(f"Conversion depuis image vers {output_format} non supportée")

        # --- Extraction texte pour les autres ---
        if ext == ".pdf":
            text = cls.pdf_to_text(input_path)
        elif ext == ".docx":
            text = cls.docx_to_text(input_path)
        elif ext == ".txt":
            text = cls.txt_to_text(input_path)
        elif ext == ".md":
            text = cls.md_to_text(input_path)
        elif ext == ".html":
            text = cls.html_to_text(input_path)
        else:
            text = ""

        # --- Conversion finale selon format ---
        if output_format == "docx":
            doc = Document()
            doc.add_paragraph(text)
            doc.save(output_path)
        elif output_format in ["txt", "md"]:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
        elif output_format == "pdf":
            cls.text_to_pdf(text, output_path)
        elif output_format == "html":
            html_text = markdown.markdown(text)
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(html_text)
        else:
            raise ValueError(f"Format de sortie non supporté: {output_format}")

        return output_path

